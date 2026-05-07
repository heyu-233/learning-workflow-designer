#!/usr/bin/env python3
"""Convert a simple Markdown learning document to DOCX.

This converter intentionally supports a conservative subset of Markdown:
headings, paragraphs, blockquote critique comments, fenced code blocks,
pipe tables, horizontal rules/page breaks, and answer lines.
"""

from __future__ import annotations

import argparse
import sys
from pathlib import Path


try:
    from docx import Document
    from docx.shared import RGBColor
except ImportError:  # pragma: no cover - exercised by environment
    Document = None
    RGBColor = None


def split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_table_separator(line: str) -> bool:
    stripped = line.strip().strip("|")
    if not stripped:
        return False
    return all(set(part.strip()) <= {"-", ":"} and "-" in part for part in stripped.split("|"))


def flush_table(document, table_rows: list[list[str]]) -> None:
    if not table_rows:
        return
    col_count = max(len(row) for row in table_rows)
    table = document.add_table(rows=len(table_rows), cols=col_count)
    table.style = "Table Grid"
    for r_idx, row in enumerate(table_rows):
        for c_idx in range(col_count):
            table.cell(r_idx, c_idx).text = row[c_idx] if c_idx < len(row) else ""


def add_red_comment(document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.font.color.rgb = RGBColor(255, 0, 0)


def convert_markdown_to_docx(input_path: Path, output_path: Path) -> None:
    if Document is None:
        raise RuntimeError("python-docx is required. Install it with: python -m pip install python-docx")

    document = Document()
    lines = input_path.read_text(encoding="utf-8").splitlines()
    in_code = False
    code_lines: list[str] = []
    table_rows: list[list[str]] = []

    def flush_code() -> None:
        nonlocal code_lines
        if code_lines:
            paragraph = document.add_paragraph()
            run = paragraph.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            code_lines = []

    def flush_pending_table() -> None:
        nonlocal table_rows
        flush_table(document, table_rows)
        table_rows = []

    for raw in lines:
        line = raw.rstrip().lstrip("\ufeff")

        if line.startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                flush_pending_table()
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if "|" in line and line.strip().startswith("|") and line.strip().endswith("|"):
            if is_table_separator(line):
                continue
            table_rows.append(split_table_row(line))
            continue
        flush_pending_table()

        if not line.strip():
            document.add_paragraph("")
            continue

        if line.strip() == "---":
            document.add_page_break()
            continue

        if line.startswith(">"):
            comment = line.lstrip(">").strip()
            if comment.startswith("批注："):
                add_red_comment(document, comment)
            else:
                document.add_paragraph(comment)
            continue

        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            text = line[level:].strip()
            document.add_heading(text, level=min(level, 4))
            continue

        document.add_paragraph(line)

    flush_code()
    flush_pending_table()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Convert simple Markdown to DOCX.")
    parser.add_argument("input", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args(argv)

    try:
        convert_markdown_to_docx(args.input, args.output)
    except Exception as exc:
        print(f"md_to_docx failed: {exc}", file=sys.stderr)
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
